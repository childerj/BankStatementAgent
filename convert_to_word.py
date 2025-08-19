#!/usr/bin/env python3
"""
Markdown to Word Converter for Bank Statement Agent Documentation
Converts the comprehensive documentation from Markdown to a downloadable Word document.
"""

import markdown
from docx import Document
from docx.shared import Inches, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.style import WD_STYLE_TYPE
from docx.oxml.shared import OxmlElement, qn
import re
import os
from datetime import datetime

def create_word_document():
    """Create a Word document from the Markdown documentation."""
    
    # Read the markdown file
    with open('BankStatementAgent_Documentation.md', 'r', encoding='utf-8') as f:
        md_content = f.read()
    
    # Create Word document
    doc = Document()
    
    # Set document properties
    doc.core_properties.title = "Bank Statement Processing Agent - Complete Documentation"
    doc.core_properties.author = "BankStatementAgent Development Team"
    doc.core_properties.subject = "Azure AI Agent Documentation"
    doc.core_properties.created = datetime.now()
    
    # Add title page
    create_title_page(doc)
    
    # Process markdown content
    process_markdown_content(doc, md_content)
    
    # Save document
    output_file = 'BankStatementAgent_Documentation.docx'
    doc.save(output_file)
    print(f"✅ Documentation exported to: {output_file}")
    print(f"📄 File size: {os.path.getsize(output_file) / 1024:.1f} KB")
    return output_file

def create_title_page(doc):
    """Create a professional title page."""
    
    # Title
    title = doc.add_heading('Bank Statement Processing Agent', 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    # Subtitle
    subtitle = doc.add_heading('Complete Technical Documentation', level=2)
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    # Add some space
    doc.add_paragraph()
    doc.add_paragraph()
    
    # Description
    desc = doc.add_paragraph()
    desc.alignment = WD_ALIGN_PARAGRAPH.CENTER
    desc.add_run("Autonomous AI-Powered Bank Statement Processing System\n").bold = True
    desc.add_run("Azure Functions • OpenAI • Document Intelligence • BAI2 Conversion")
    
    # Add more space
    doc.add_paragraph()
    doc.add_paragraph()
    
    # Version info
    version_para = doc.add_paragraph()
    version_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    version_para.add_run("Version: 1.0\n")
    version_para.add_run(f"Generated: {datetime.now().strftime('%B %d, %Y')}\n")
    version_para.add_run("Development Team: BankStatementAgent Project")
    
    # Page break
    doc.add_page_break()

def process_markdown_content(doc, md_content):
    """Process markdown content and add to Word document."""
    
    lines = md_content.split('\n')
    i = 0
    
    while i < len(lines):
        line = lines[i].strip()
        
        if not line:
            i += 1
            continue
            
        # Skip title page content (already added)
        if line.startswith('# Bank Statement Processing Agent'):
            i += 1
            continue
            
        # Process headers
        if line.startswith('#'):
            level = len(line) - len(line.lstrip('#'))
            title = line.lstrip('# ').strip()
            
            # Skip TOC
            if title == "Table of Contents":
                # Skip until next header
                i += 1
                while i < len(lines) and not lines[i].strip().startswith('#'):
                    i += 1
                continue
                
            # Add header
            if level <= 3:  # Only process levels 1-3
                heading = doc.add_heading(title, level)
                if level == 1:
                    heading.alignment = WD_ALIGN_PARAGRAPH.LEFT
            
        # Process code blocks
        elif line.startswith('```'):
            i, code_block = process_code_block(lines, i)
            if code_block:
                add_code_block(doc, code_block)
            continue
            
        # Process mermaid diagrams (convert to text description)
        elif line.startswith('```mermaid'):
            i, diagram_content = process_mermaid_diagram(lines, i)
            if diagram_content:
                add_diagram_placeholder(doc, diagram_content)
            continue
            
        # Process tables
        elif '|' in line and '---' not in line:
            i, table_data = process_table(lines, i)
            if table_data:
                add_table(doc, table_data)
            continue
            
        # Process regular paragraphs
        else:
            paragraph = process_paragraph(line)
            if paragraph:
                p = doc.add_paragraph(paragraph)
                
        i += 1

def process_code_block(lines, start_idx):
    """Process code blocks."""
    i = start_idx + 1
    code_lines = []
    language = lines[start_idx].replace('```', '').strip()
    
    while i < len(lines) and not lines[i].strip().startswith('```'):
        code_lines.append(lines[i])
        i += 1
    
    return i + 1, {
        'language': language,
        'content': '\n'.join(code_lines)
    }

def process_mermaid_diagram(lines, start_idx):
    """Process mermaid diagrams."""
    i = start_idx + 1
    diagram_lines = []
    
    while i < len(lines) and not lines[i].strip().startswith('```'):
        diagram_lines.append(lines[i])
        i += 1
    
    return i + 1, '\n'.join(diagram_lines)

def process_table(lines, start_idx):
    """Process markdown tables."""
    table_rows = []
    i = start_idx
    
    while i < len(lines) and '|' in lines[i]:
        row = lines[i].strip()
        if '---' not in row:  # Skip separator rows
            cells = [cell.strip() for cell in row.split('|')[1:-1]]  # Remove empty first/last
            table_rows.append(cells)
        i += 1
    
    return i, table_rows

def process_paragraph(line):
    """Process regular paragraph text."""
    # Remove markdown formatting for Word
    text = line
    text = re.sub(r'\*\*(.*?)\*\*', r'\1', text)  # Bold
    text = re.sub(r'\*(.*?)\*', r'\1', text)      # Italic
    text = re.sub(r'`(.*?)`', r'\1', text)        # Code
    text = re.sub(r'\[(.*?)\]\(.*?\)', r'\1', text)  # Links
    
    return text

def add_code_block(doc, code_block):
    """Add a code block to the document."""
    # Add language label if specified
    if code_block['language']:
        label = doc.add_paragraph()
        label.add_run(f"Code ({code_block['language'].upper()}):").bold = True
    
    # Add code content
    code_para = doc.add_paragraph()
    code_run = code_para.add_run(code_block['content'])
    code_run.font.name = 'Courier New'
    code_run.font.size = Pt(9)
    
    # Add background shading
    shading_elm = OxmlElement('w:shd')
    shading_elm.set(qn('w:fill'), 'F8F8F8')
    code_para._element.get_or_add_pPr().append(shading_elm)

def add_diagram_placeholder(doc, diagram_content):
    """Add a placeholder for mermaid diagrams."""
    # Add diagram heading
    diagram_para = doc.add_paragraph()
    diagram_para.add_run("Diagram: ").bold = True
    diagram_para.add_run("(Architecture/Flow Diagram)")
    
    # Add simplified description
    desc_para = doc.add_paragraph()
    desc_para.add_run("Note: This section contains a system architecture or flow diagram. ")
    desc_para.add_run("Please refer to the original markdown file or web version for visual diagrams.")
    
    # Add a simple text representation
    if 'graph' in diagram_content.lower():
        desc_para = doc.add_paragraph()
        desc_para.add_run("Diagram Type: System Architecture Flow")
    elif 'sequence' in diagram_content.lower():
        desc_para = doc.add_paragraph()
        desc_para.add_run("Diagram Type: Sequence/Process Flow")

def add_table(doc, table_data):
    """Add a table to the document."""
    if not table_data:
        return
        
    table = doc.add_table(rows=len(table_data), cols=len(table_data[0]))
    table.style = 'Table Grid'
    
    for row_idx, row_data in enumerate(table_data):
        for col_idx, cell_data in enumerate(row_data):
            cell = table.cell(row_idx, col_idx)
            cell.text = cell_data
            
            # Make header row bold
            if row_idx == 0:
                for paragraph in cell.paragraphs:
                    for run in paragraph.runs:
                        run.bold = True

def main():
    """Main function to convert documentation."""
    print("🔄 Converting Bank Statement Agent Documentation to Word...")
    print("📝 Processing markdown content...")
    
    try:
        output_file = create_word_document()
        print(f"✅ SUCCESS: Documentation converted to {output_file}")
        print("📁 The Word document is ready for download!")
        
        # Print file info
        file_path = os.path.abspath(output_file)
        print(f"📍 Full path: {file_path}")
        
    except Exception as e:
        print(f"❌ ERROR: Failed to convert documentation: {e}")
        raise

if __name__ == "__main__":
    main()
